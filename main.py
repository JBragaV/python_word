from io import StringIO

from docx import Document


def iniciar(caminho_arquivo=None):
    if caminho_arquivo is not None:
        arquivo = open(caminho_arquivo, 'rb')
        documento = Document(arquivo)
        # documento.add_heading('Esse é o novo Heading 0 Formulário de Trocas do mês', 0)
        # documento.add_heading('Heading 1', 1)
        # documento.add_heading('Heading 2', 2)
        # documento.add_heading('Heading 3', 3)
        # documento.add_heading('Heading 4', 4)
        # documento.add_heading('Heading 5', 5)
        # documento.add_heading('Heading 6', 6)
        # documento.add_heading('Heading 7', 7)
        # documento.add_heading('Heading 8', 8)
        # documento.add_heading('Heading 9', 9)
        cursos = [
            ['Felipe', 'Python', 5, 'Jocimar', 69, "18/11/1982"],
            ['Alisson', 'Android', 4, 'Dadada', 69, "06/05/1988"],
            ['Haynes', 'Python', 4, 'Manuzinha linda', 9, "20/06/2013"]
            # (Instrutor, Categoria, Quantidade)
        ]
        tabela = documento.add_table(rows=1, cols=6)
        # Adicionamos apenas 1 linha (o titulo)
        titulo_tabela = tabela.rows[0].cells
        # Pegamos as células da tabela
        titulo_tabela[0].text = "Nome do Instrutor"
        titulo_tabela[1].text = "Categoria do Curso"
        titulo_tabela[2].text = "Quantidade de Cursos"
        titulo_tabela[3].text = "Nome de Aluno"
        titulo_tabela[4].text = "Quantidade qualquer"
        titulo_tabela[5].text = "Data"
        for nome, categoria, quantidade, aluno, numero, data in cursos:
            linha = tabela.add_row().cells
            linha[0].text = nome
            linha[1].text = categoria
            linha[2].text = str(quantidade)
            linha[3].text = aluno
            linha[4].text = str(numero)
            linha[5].text = data
        documento.save(caminho_arquivo)
    else:
        documento = Document()
        documento.add_heading('Heading 0 Formulário de Trocas do mês', 0)
        documento.add_heading('Heading 1', 1)
        documento.add_heading('Heading 2', 2)
        documento.add_heading('Heading 3', 3)
        documento.add_heading('Heading 4', 4)
        documento.add_heading('Heading 5', 5)
        documento.add_heading('Heading 6', 6)
        documento.add_heading('Heading 7', 7)
        documento.add_heading('Heading 8', 8)
        documento.add_heading('Heading 9', 9)
        cursos = [
            ['Felipe', 'Python', 5, 'Jocimar', 69, "18/11/1982"],
            ['Alisson', 'Android', 4, 'Dadada', 69, "06/05/1988"],
            ['Haynes', 'Python', 4, 'Manuzinha linda', 9, "20/06/2013"]
            # (Instrutor, Categoria, Quantidade)
        ]
        tabela = documento.add_table(rows=1, cols=6)
        # Adicionamos apenas 1 linha (o titulo)
        titulo_tabela = tabela.rows[0].cells
        # Pegamos as células da tabela
        titulo_tabela[0].text = "Nome do Instrutor"
        titulo_tabela[1].text = "Categoria do Curso"
        titulo_tabela[2].text = "Quantidade de Cursos"
        titulo_tabela[3].text = "Nome de Aluno"
        titulo_tabela[4].text = "Quantidade qualquer"
        titulo_tabela[5].text = "Data"
        for nome, categoria, quantidade, aluno, numero, data in cursos:
            linha = tabela.add_row().cells
            linha[0].text = nome
            linha[1].text = categoria
            linha[2].text = str(quantidade)
            linha[3].text = aluno
            linha[4].text = str(numero)
            linha[5].text = data
        documento.save("arquivo_novo_teste.docx")


if __name__ == '__main__':
    # iniciar(r"C:\Users\jbrag\PycharmProjects\python_word\arquivo_word_teste.docx")
    iniciar("arquivo_novo_teste.docx")
    # iniciar()
