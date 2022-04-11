from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from datetime import datetime
import excel as ex

LISTA_MESES = ["Janeiro", "Fevereiro", "Março", "Abril",
               "Maio", "Junho", "Julho", "Agosto",
               "Setembro", "Outubro", "Novembro", "Dezembro"]


def iniciar(caminho_arquivo=None):
    mes_sistema = LISTA_MESES[datetime.now().month-1]
    if caminho_arquivo is not None:
        print(__name__)
        trocas = ex.abrir_excel()
        print(trocas)
        arquivo = open(caminho_arquivo, 'rb')
        documento = Document(arquivo)
        paragrafo_estilo = documento.add_paragraph("")
        # paragrafo_estilo.style = documento.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
        titulo = documento.add_heading(f"Trocas de {mes_sistema}", 0)
        font = titulo.style.font
        font.name = 'Time New Roman'
        font.size = Pt(24)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        """trocas = [
            ['Felipe', '24/06/2024', 'A', 'Jocimar', "18/11/1982", 'D'],
            ['Jocimar', '18/11/1982', 'B', 'Dadada', "06/05/1988", 'A'],
            ['Haynes', '24/11/2020', 'C', 'Manuela', "20/06/2013", 'B'],
            ['Felipe', '24/06/2024', 'A', 'Jocimar', "18/11/1982", 'D'],
            # (Instrutor, Categoria, Quantidade)
        ]"""
        colunas_tabela = 6
        tabela = documento.add_table(rows=1, cols=colunas_tabela)
        # tabela.rows[0].cells[0].text =
        # a = tabela.cell(0, 0)
        # b = tabela.cell(0, 1)
        # c = tabela.cell(0, 2)
        # d = tabela.cell(0, 3)
        # e = tabela.cell(0, 4)
        # f = tabela.cell(0, 5)
        # A = a.merge(b)
        # B = A.merge(c)
        # C = B.merge(d)
        # D = C.merge(e)
        # E = D.merge(f)
        # # Fim
        # Inicio da formatação das células do solicitante
        linha_dados = 0
        tabela.rows[linha_dados].cells[0].text = "Solicitante"
        a = tabela.cell(linha_dados, 0)
        b = tabela.cell(linha_dados, 1)
        c = tabela.cell(linha_dados, 2)
        A = a.merge(b)
        B = A.merge(c)
        # Fim
        # Inicio da formatação das células do solicitado
        tabela.rows[linha_dados].cells[3].text = "Solicitado"
        d = tabela.cell(linha_dados, 3)
        e = tabela.cell(linha_dados, 4)
        f = tabela.cell(linha_dados, 5)
        C = d.merge(e)
        D = C.merge(f)
        # Fim

        # Inicio da incerção do cabeçalho dos campos necessários. indice 3
        linha = tabela.add_row().cells
        campos_trocas = ["Nome", "Data Serviço", "Turno", "Nome", "Data Serviço", "Turno"]
        for indice, dados in enumerate(campos_trocas):
            linha[indice].text = dados

        for nome, categoria, quantidade, aluno, numero, data in trocas:
            linha = tabela.add_row().cells
            linha[0].text = nome
            linha[1].text = categoria
            linha[2].text = str(quantidade)
            linha[3].text = aluno
            linha[4].text = str(numero)
            linha[5].text = data
        # Adicionamos apenas 1 linha (o titulo)
        # titulo_tabela = tabela.rows[0].cells
        # # Pegamos as células da tabela
        # titulo_tabela[0].text = "Nome do Instrutor"
        # titulo_tabela[1].text = "Categoria do Curso"
        # titulo_tabela[2].text = "Quantidade de Cursos"
        # titulo_tabela[3].text = "Nome de Aluno"
        # titulo_tabela[4].text = "Quantidade qualquer"
        # titulo_tabela[5].text = "Data"
        # for nome, categoria, quantidade, aluno, numero, data in trocas:
        #     linha = tabela.add_row().cells
        #     linha[0].text = nome
        #     linha[1].text = categoria
        #     linha[2].text = str(quantidade)
        #     linha[3].text = aluno
        #     linha[4].text = str(numero)
        #     linha[5].text = data
        #
        # print(tabela.rows[2].cells[3].text)
        documento.save(f"troca_{mes_sistema}.docx")
        return
    print("Deu um erro aí!!!")



